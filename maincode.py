import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix, ConfusionMatrixDisplay
import joblib
from docx import Document
from docx.shared import Inches

TRAIN_CSV = "train.csv"
TEST_CSV = "test.csv"
OUTPUT_MODEL = "rf_instagram_fake_model.joblib"
PREDICTIONS_CSV = "test_predictions.csv"
REPORT_FILE = "Instagram_Project_Report.docx"

df_train = pd.read_csv(TRAIN_CSV)
df_test = pd.read_csv(TEST_CSV)

# Features and target
feature_cols = [
    'profile pic', 'nums/length username', 'fullname words',
    'nums/length fullname', 'name==username', 'description length',
    'external URL', 'private', '#posts', '#followers', '#follows'
]
X = df_train[feature_cols].copy()
y = df_train['fake']

# Feature Engineering
X['follower_following_ratio'] = X['#followers'] / (X['#follows'] + 1)

# Train/validation split
X_train, X_val, y_train, y_val = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)

# Scaling
scaler = StandardScaler()
X_train_s = scaler.fit_transform(X_train)
X_val_s = scaler.transform(X_val)

# Models

# Decision Tree
dt = DecisionTreeClassifier(random_state=42)
dt.fit(X_train_s, y_train)
y_pred_dt = dt.predict(X_val_s)
dt_acc = accuracy_score(y_val, y_pred_dt)

# Random Forest
rf = RandomForestClassifier(n_estimators=200, random_state=42, n_jobs=-1)
rf.fit(X_train_s, y_train)
y_pred_rf = rf.predict(X_val_s)
rf_acc = accuracy_score(y_val, y_pred_rf)
rf_report = classification_report(y_val, y_pred_rf)
rf_conf = confusion_matrix(y_val, y_pred_rf)

# Feature importance
importances = rf.feature_importances_
feat_importance = pd.Series(importances, index=X.columns).sort_values(ascending=False)

# Save model & predictions
joblib.dump({'model': rf, 'scaler': scaler, 'features': list(X.columns)}, OUTPUT_MODEL)

X_test = df_test[feature_cols].copy()
X_test['follower_following_ratio'] = X_test['#followers'] / (X_test['#follows'] + 1)
X_test_s = scaler.transform(X_test)
test_preds = rf.predict(X_test_s)
df_test['pred_fake'] = test_preds
df_test.to_csv(PREDICTIONS_CSV, index=False)

# Save Plots as Images
# Confusion matrix heatmap
plt.figure(figsize=(5, 4))
sns.heatmap(rf_conf, annot=True, fmt="d", cmap="Blues",
            xticklabels=["Genuine", "Fake"],
            yticklabels=["Genuine", "Fake"])
plt.title("Confusion Matrix - Random Forest")
plt.savefig("confusion_matrix.png")
plt.close()

# Feature importance bar chart
plt.figure(figsize=(7, 5))
sns.barplot(x=feat_importance.values[:10], y=feat_importance.index[:10], palette="viridis")
plt.title("Top 10 Feature Importances - Random Forest")
plt.tight_layout()
plt.savefig("feature_importance.png")
plt.close()

# Generate Word Report
def create_project_report(filename=REPORT_FILE):
    doc = Document()
    doc.add_heading("Instagram Fake vs Genuine Account Detection", 0)

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "This project detects fake Instagram accounts using machine learning. "
        "It applies data analysis, feature engineering, and classification models "
        "to distinguish between genuine and fake accounts."
    )

    doc.add_heading("2. Dataset Details", level=1)
    doc.add_paragraph(
        f"Training dataset: {df_train.shape[0]} rows, {df_train.shape[1]} columns\n"
        f"Testing dataset: {df_test.shape[0]} rows, {df_test.shape[1]} columns\n\n"
        "Features:\n"
        "- Profile picture (binary)\n"
        "- Username length and digits\n"
        "- Full name length/words\n"
        "- Bio/description length\n"
        "- Private/Public status\n"
        "- Posts, Followers, Following\n"
        "- Target: Fake (1) or Genuine (0)"
    )

    doc.add_heading("3. Program Features", level=1)
    doc.add_paragraph(
        "- Loads training and test datasets\n"
        "- Performs feature engineering (follower/following ratio)\n"
        "- Trains Decision Tree and Random Forest models\n"
        "- Evaluates models with accuracy, precision, recall, F1-score\n"
        "- Saves trained model with Joblib\n"
        "- Exports predictions on test data"
    )

    doc.add_heading("4. Results", level=1)
    doc.add_paragraph(
        f"Decision Tree Accuracy: {dt_acc:.2%}\n"
        f"Random Forest Accuracy: {rf_acc:.2%}\n\n"
        f"Confusion Matrix (Random Forest):\n{rf_conf}\n\n"
        f"Classification Report:\n{rf_report}\n"
    )

    # Insert Confusion Matrix plot
    doc.add_paragraph("Confusion Matrix Visualization:")
    doc.add_picture("confusion_matrix.png", width=Inches(4.5))

    # Insert Feature Importance plot
    doc.add_paragraph("Top Feature Importances:")
    doc.add_picture("feature_importance.png", width=Inches(4.5))

    doc.add_heading("5. Future Improvements", level=1)
    doc.add_paragraph(
        "- Try XGBoost or Gradient Boosting\n"
        "- Handle imbalanced data with SMOTE\n"
        "- Deploy as a web app for real-time detection\n"
        "- Add more social graph/network features"
    )

    doc.save(filename)
    print(f"Report saved as {filename}")

# Generate report
create_project_report()